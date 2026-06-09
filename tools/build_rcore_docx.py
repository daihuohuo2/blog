from __future__ import annotations

import re
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


MD_PATH = Path("source/_posts/rcore-ch3-ch8-lab-review.md")
OUT_PATH = Path("build/rcore-ch3-ch8-lab-review.docx")


def set_east_asian_font(run, font_name: str) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_paragraph_border(paragraph, color="D7DEE8", size="6", space="4") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)


def set_paragraph_shading(paragraph, fill="F6F8FA") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def clean_inline(text: str) -> str:
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1（\2）", text)
    text = text.replace("&lt;", "<").replace("&gt;", ">").replace("&amp;", "&")
    return text.strip()


def split_table_row(line: str) -> list[str]:
    line = line.strip().strip("|")
    return [clean_inline(cell.strip()) for cell in line.split("|")]


def is_table_separator(line: str) -> bool:
    stripped = line.strip()
    if not stripped.startswith("|"):
        return False
    cells = [c.strip() for c in stripped.strip("|").split("|")]
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", c or "") for c in cells)


def add_normal_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.space_after = Pt(6)
    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = p.add_run(part[1:-1])
            set_east_asian_font(run, "Consolas")
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(51, 65, 85)
        else:
            run = p.add_run(clean_inline(part))
            set_east_asian_font(run, "Microsoft YaHei")


def add_code_block(doc: Document, code: str, language: str = "") -> None:
    if language:
        label = doc.add_paragraph(style="Normal")
        label.paragraph_format.space_before = Pt(4)
        label.paragraph_format.space_after = Pt(2)
        r = label.add_run(f"代码 / 命令（{language}）")
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(71, 85, 105)
        set_east_asian_font(r, "Microsoft YaHei")
    for raw in code.rstrip("\n").splitlines() or [""]:
        p = doc.add_paragraph(style="Code Block")
        p.paragraph_format.left_indent = Cm(0.25)
        p.paragraph_format.right_indent = Cm(0.15)
        p.paragraph_format.space_after = Pt(0)
        set_paragraph_shading(p)
        run = p.add_run(raw if raw else " ")
        set_east_asian_font(run, "Consolas")
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(30, 41, 59)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)


def add_markdown_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = True
    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if r_idx == 0:
                set_cell_shading(cell, "EEF2F7")
            text = row[c_idx] if c_idx < len(row) else ""
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(text) < 24 else WD_ALIGN_PARAGRAPH.LEFT
            run = paragraph.add_run(text)
            set_east_asian_font(run, "Microsoft YaHei")
            run.font.size = Pt(9)
            if r_idx == 0:
                run.bold = True
                run.font.color.rgb = RGBColor(30, 41, 59)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)

    for name, size, color in [
        ("Title", 22, "1F4E79"),
        ("Heading 1", 16, "1F4E79"),
        ("Heading 2", 13, "334155"),
        ("Heading 3", 11.5, "475569"),
    ]:
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(10 if name != "Title" else 0)
        style.paragraph_format.space_after = Pt(5)

    if "Code Block" not in [s.name for s in styles]:
        code_style = styles.add_style("Code Block", 1)
    else:
        code_style = styles["Code Block"]
    code_style.font.name = "Consolas"
    code_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    code_style.font.size = Pt(8.5)
    code_style.paragraph_format.line_spacing = 1.0

    header = section.header
    hp = header.paragraphs[0]
    hp.text = ""
    hr = hp.add_run("rCore ch3-ch8 实验复盘")
    set_east_asian_font(hr, "Microsoft YaHei")
    hr.font.size = Pt(9)
    hr.font.color.rgb = RGBColor(100, 116, 139)
    set_paragraph_border(hp)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fr = fp.add_run("由 Markdown 复盘笔记生成")
    set_east_asian_font(fr, "Microsoft YaHei")
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor(148, 163, 184)


def build_docx() -> None:
    source = MD_PATH.read_text(encoding="utf-8")
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("rCore ch3-ch8 实验复盘笔记")
    set_east_asian_font(run, "Microsoft YaHei")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = subtitle.add_run(f"根据 source/_posts/rcore-ch3-ch8-lab-review.md 生成 · {datetime.now():%Y-%m-%d}")
    set_east_asian_font(sr, "Microsoft YaHei")
    sr.font.size = Pt(10)
    sr.font.color.rgb = RGBColor(100, 116, 139)

    doc.add_paragraph()

    lines = source.splitlines()
    i = 0
    in_code = False
    code_lang = ""
    code_lines: list[str] = []

    while i < len(lines):
        line = lines[i]

        if line.startswith("```"):
            if in_code:
                add_code_block(doc, "\n".join(code_lines), code_lang)
                code_lines = []
                code_lang = ""
                in_code = False
            else:
                in_code = True
                code_lang = line.strip("`").strip() or "text"
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if not line.strip():
            i += 1
            continue

        if line.strip().startswith("|") and i + 1 < len(lines) and is_table_separator(lines[i + 1]):
            rows = [split_table_row(line)]
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(split_table_row(lines[i]))
                i += 1
            add_markdown_table(doc, rows)
            continue

        heading = re.match(r"^(#{1,6})\s+(.*)$", line)
        if heading:
            level = len(heading.group(1))
            text = clean_inline(heading.group(2))
            if level == 1:
                if doc.paragraphs and doc.paragraphs[-1].text.strip():
                    doc.add_section(WD_SECTION.NEW_PAGE)
                p = doc.add_paragraph(text, style="Heading 1")
            elif level == 2:
                p = doc.add_paragraph(text, style="Heading 1")
            elif level == 3:
                p = doc.add_paragraph(text, style="Heading 2")
            else:
                p = doc.add_paragraph(text, style="Heading 3")
            p.paragraph_format.keep_with_next = True
            i += 1
            continue

        if line.startswith(">"):
            p = doc.add_paragraph(style="Intense Quote")
            r = p.add_run(clean_inline(line.lstrip("> ").strip()))
            set_east_asian_font(r, "Microsoft YaHei")
            r.font.size = Pt(10)
            i += 1
            continue

        bullet = re.match(r"^\s*[-*]\s+(.*)$", line)
        if bullet:
            p = doc.add_paragraph(style="List Bullet")
            r = p.add_run(clean_inline(bullet.group(1)))
            set_east_asian_font(r, "Microsoft YaHei")
            i += 1
            continue

        numbered = re.match(r"^\s*\d+[.)]\s+(.*)$", line)
        if numbered:
            p = doc.add_paragraph(style="List Number")
            r = p.add_run(clean_inline(numbered.group(1)))
            set_east_asian_font(r, "Microsoft YaHei")
            i += 1
            continue

        add_normal_paragraph(doc, line)
        i += 1

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_PATH)
    print(OUT_PATH.resolve())


if __name__ == "__main__":
    try:
        build_docx()
    except Exception as exc:
        print(f"failed: {exc}", file=sys.stderr)
        raise
