from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "HTTP代理服务器实验报告.docx"


def set_run_font(run, east_asia="宋体", ascii_font="Arial", size=None, bold=None):
    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def add_paragraph(doc, text="", style=None):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    set_run_font(run, size=11)
    return p


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    set_run_font(run, size=10.5, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_heading(doc, text, level=1):
    p = doc.add_heading(level=level)
    run = p.add_run(text)
    set_run_font(run, east_asia="黑体", ascii_font="Arial", size=16 if level == 1 else 13, bold=True)
    run.font.color.rgb = RGBColor(31, 78, 121)
    return p


def add_code_block(doc, code):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade_cell(cell, "F6F8FA")
    cell.text = ""
    for line in code.strip("\n").splitlines():
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        set_run_font(run, east_asia="Consolas", ascii_font="Consolas", size=9)


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("HTTP 代理服务器实验报告")
    set_run_font(run, east_asia="黑体", ascii_font="Arial", size=20, bold=True)
    run.font.color.rgb = RGBColor(31, 78, 121)

    meta = doc.add_table(rows=4, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta.style = "Table Grid"
    rows = [
        ("实验题目", "附录3-A类实验第2题：搭建 HTTP 代理"),
        ("开发语言", "C 语言，Berkeley Socket API"),
        ("运行环境", "Linux / WSL，gcc，Makefile"),
        ("程序文件", "proxy.c、Makefile"),
    ]
    for row, (k, v) in zip(meta.rows, rows):
        shade_cell(row.cells[0], "D9EAF7")
        set_cell_text(row.cells[0], k, bold=True)
        set_cell_text(row.cells[1], v)

    add_heading(doc, "一、实验目的")
    for item in [
        "掌握 HTTP/1.0 请求与响应的基本格式。",
        "熟悉 Berkeley 套接字编程流程，包括 socket、bind、listen、accept、connect、send 和 recv。",
        "实现一个能够在客户端与 Web 服务器之间转发数据的 HTTP 代理服务器。",
        "使用 fork() 为多个客户端连接创建子进程，实现并发处理。",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_run_font(run, size=11)

    add_heading(doc, "二、功能设计")
    add_paragraph(
        doc,
        "程序启动时通过命令行参数指定监听端口，主进程负责接收客户端连接。"
        "每当有新的客户端连接到来时，父进程调用 fork() 创建子进程，子进程负责解析请求、连接远程服务器、转发请求并回传响应。"
    )
    design = doc.add_table(rows=1, cols=3)
    design.style = "Table Grid"
    design.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, text in enumerate(["模块", "作用", "实现要点"]):
        shade_cell(design.rows[0].cells[i], "D9EAF7")
        set_cell_text(design.rows[0].cells[i], text, bold=True)
    for values in [
        ("监听模块", "等待客户端连接", "getaddrinfo、socket、bind、listen、accept"),
        ("并发处理", "让多个客户端同时访问代理", "每个连接 fork 一个子进程，父进程继续 accept"),
        ("请求解析", "提取方法、URL、主机、端口、路径", "支持绝对 URL，也兼容带 Host 头的普通路径"),
        ("请求转发", "连接远程 Web 服务器", "只实现 GET，其他方法返回 501"),
        ("响应回传", "把服务器响应原样发回客户端", "循环 recv，再完整 send 给客户端"),
    ]:
        row = design.add_row()
        for i, text in enumerate(values):
            set_cell_text(row.cells[i], text)

    add_heading(doc, "三、关键处理逻辑")
    add_paragraph(doc, "代理收到合法 GET 请求后，将请求行改写为面向源站服务器的 HTTP/1.0 请求，并补充关闭连接相关头部：")
    add_code_block(
        doc,
        """
GET /path HTTP/1.0
Host: example.com
Connection: close
Proxy-Connection: close
""",
    )
    add_paragraph(
        doc,
        "若请求方法不是 GET，程序直接返回 501 Not Implemented；若 URL 或 Host 信息无效，则返回 400 Bad Request；"
        "若无法连接远程服务器，则返回 502 Bad Gateway。"
    )

    add_heading(doc, "四、编译与运行")
    add_code_block(
        doc,
        """
make
./proxy 8080
""",
    )
    add_paragraph(doc, "使用 telnet 或浏览器代理设置进行测试。telnet 示例：")
    add_code_block(
        doc,
        """
telnet 127.0.0.1 8080
GET http://example.com/ HTTP/1.0

""",
    )

    add_heading(doc, "五、测试结果")
    result = doc.add_table(rows=1, cols=4)
    result.style = "Table Grid"
    result.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, text in enumerate(["测试项", "输入", "期望结果", "实际结果"]):
        shade_cell(result.rows[0].cells[i], "D9EAF7")
        set_cell_text(result.rows[0].cells[i], text, bold=True)
    for values in [
        ("GET 转发", "GET http://127.0.0.1:18080/ HTTP/1.0", "返回源站页面内容", "HTTP/1.0 200 OK，正文 proxy-ok"),
        ("非 GET 方法", "POST http://127.0.0.1:18080/ HTTP/1.0", "返回 501", "HTTP/1.0 501 Not Implemented"),
    ]:
        row = result.add_row()
        for i, text in enumerate(values):
            set_cell_text(row.cells[i], text)

    add_heading(doc, "六、结论")
    add_paragraph(
        doc,
        "本实验完成了一个基础 HTTP 代理服务器。程序能够监听指定端口，解析客户端 HTTP 请求，"
        "将 GET 请求转发到远程服务器，并把响应数据原样返回客户端；同时按照实验要求对未实现的请求方法返回 501。"
    )

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
