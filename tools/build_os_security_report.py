from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


OUT = r"C:\Users\20694\Desktop\myblog\操作系统安全调研报告_CVE-2023-28252.docx"


def fmt(run, size=12, bold=False):
    run.font.name = "SimSun"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)


def p(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.first_line_indent = Cm(0.74)
    para.paragraph_format.line_spacing = 1.25
    para.paragraph_format.space_after = Pt(4)
    run = para.add_run(text)
    fmt(run, 12)


def h(doc, text, level=1):
    para = doc.add_heading(level=level)
    run = para.add_run(text)
    fmt(run, 14 if level == 1 else 12, True)
    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after = Pt(4)


def bullet(doc, text):
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.left_indent = Cm(0.75)
    para.paragraph_format.space_after = Pt(2)
    run = para.add_run(text)
    fmt(run, 12)


def table_style(table):
    table.style = "Table Grid"
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for para in cell.paragraphs:
                para.paragraph_format.space_after = Pt(0)
                for run in para.runs:
                    fmt(run, 11)


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.9)
section.bottom_margin = Inches(0.9)
section.left_margin = Inches(0.9)
section.right_margin = Inches(0.9)

styles = doc.styles
for name in ["Normal", "Heading 1", "Heading 2", "List Bullet"]:
    style = styles[name]
    style.font.name = "SimSun"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.color.rgb = RGBColor(0, 0, 0)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("操作系统安全调研报告")
fmt(r, 18, True)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("案例：Windows CLFS 提权漏洞 CVE-2023-28252")
fmt(r, 12, False)

h(doc, "一、案例简介")
p(doc, "本文选择 Windows Common Log File System（CLFS）本地提权漏洞 CVE-2023-28252 作为案例。该漏洞在 2023 年被发现，并且已经被 Nokoyawa 勒索软件相关攻击者真实利用。它不是网页漏洞，而是 Windows 操作系统底层驱动 clfs.sys 中的问题。")
p(doc, "CLFS 是 Windows 的通用日志系统，用来处理系统或程序产生的日志文件。攻击者可以构造恶意的 BLF 日志文件，让系统驱动在解析时发生越界写，最后把普通用户权限提升到 SYSTEM 权限。")

h(doc, "二、攻击过程和原理")
bullet(doc, "攻击者先通过其他方式进入 Windows 服务器，获得普通用户权限。")
bullet(doc, "然后在目标机器上运行提权利用程序，并准备被修改过的 BLF 日志文件。")
bullet(doc, "clfs.sys 驱动解析该日志文件时发生越界写，导致内核内存被破坏。")
bullet(doc, "攻击者借此把普通权限提升到 SYSTEM 权限。")
bullet(doc, "提权后，攻击者可以继续窃取凭据、关闭安全软件，甚至投放 Nokoyawa 勒索软件。")
p(doc, "简单来说，这个漏洞的关键点是：普通用户能控制一部分日志文件内容，而系统驱动在处理这些内容时没有检查好边界，于是把数据写到了不该写的位置。由于驱动运行在内核态，所以后果比普通软件漏洞更严重。")

h(doc, "三、影响范围")
p(doc, "该漏洞影响多个 Windows 版本，包括 Windows 10、Windows 11 以及 Windows Server。它本身不是远程直接打入系统的漏洞，但如果攻击者已经能在机器上运行程序，就可以利用它扩大权限。")
p(doc, "在真实攻击中，攻击者利用它部署 Nokoyawa 勒索软件。拿到 SYSTEM 权限后，攻击者可以读取敏感文件、获取密码哈希、创建系统服务、删除备份，对企业影响很大。")

h(doc, "四、涉及的操作系统知识点")
bullet(doc, "用户态和内核态：普通程序在用户态运行，系统驱动在内核态运行。")
bullet(doc, "权限管理：Windows 中 SYSTEM 权限很高，很多系统资源都可以访问。")
bullet(doc, "驱动程序：clfs.sys 是系统驱动，驱动漏洞常常会导致本地提权。")
bullet(doc, "文件系统和日志：CLFS 负责处理日志文件，恶意 BLF 文件成为触发漏洞的入口。")
bullet(doc, "内存安全：越界写说明程序写到了不该写的内存区域。")

h(doc, "五、同类案例")
table = doc.add_table(rows=1, cols=3)
table.cell(0, 0).text = "案例"
table.cell(0, 1).text = "系统"
table.cell(0, 2).text = "相似点"
for a, b, c in [
    ("CVE-2022-37969", "Windows CLFS", "同样是 CLFS 本地提权漏洞"),
    ("CVE-2021-4034", "Linux Polkit", "普通用户可提权到 root"),
    ("CVE-2023-0386", "Linux OverlayFS", "文件系统相关问题导致提权"),
]:
    cells = table.add_row().cells
    cells[0].text = a
    cells[1].text = b
    cells[2].text = c
table_style(table)

h(doc, "六、防御和加固")
bullet(doc, "及时安装 Microsoft 安全补丁，尤其是已经被真实利用的漏洞。")
bullet(doc, "服务器尽量不要使用普通账号长期登录，也不要随便给管理员权限。")
bullet(doc, "开启 Defender、EDR 等终端防护，监控异常进程和可疑日志文件。")
bullet(doc, "限制未知程序运行，减少攻击者执行提权工具的机会。")
bullet(doc, "做好离线备份，防止勒索软件加密生产数据后无法恢复。")

h(doc, "七、简单总结")
p(doc, "这个案例说明，操作系统安全不只是防网页攻击。攻击者进入系统后，还会继续寻找本地提权漏洞，把普通权限变成最高权限。CLFS 原本只是日志系统，但因为驱动解析文件时存在问题，就被攻击者利用成了提权工具。")
p(doc, "所以防御时不能只靠一个办法。系统补丁、权限控制、终端检测和备份恢复都很重要。特别是这种已经被真实攻击利用的漏洞，发现后应该尽快修复。")

h(doc, "参考资料")
refs = [
    "Microsoft Security Response Center：CVE-2023-28252，https://msrc.microsoft.com/update-guide/vulnerability/CVE-2023-28252",
    "Kaspersky：CVE-2023-28252 zero-day vulnerability in CLFS，https://www.kaspersky.com/blog/nokoyawa-zero-day-exploit/47788/",
    "Kaspersky Securelist：Nokoyawa ransomware attacks with Windows zero-day，https://securelist.com/nokoyawa-ransomware-attacks-with-windows-zero-day/109483/",
    "NVD：CVE-2023-28252，https://nvd.nist.gov/vuln/detail/CVE-2023-28252",
    "Microsoft Learn：Introduction to the Common Log File System，https://learn.microsoft.com/en-us/windows-hardware/drivers/kernel/introduction-to-the-common-log-file-system",
]
for item in refs:
    p(doc, item)

doc.save(OUT)
print(OUT)
